import pandas as pd
import json
from copy import copy
from openpyxl import load_workbook
from TaiwanLottery import TaiwanLotteryCrawler
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
from collections import Counter
from scratch_crawler import create_scratch_json
from star_lottery import get_3star, get_4star
# =====================================
# 日期格式
# =====================================

def roc_date(date_str):
    d = pd.to_datetime(date_str)
    return f"{d.year - 1911}.{d.month:02d}.{d.day:02d}"


def short_date(date_str):

    week_map = {
        0: "一",
        1: "二",
        2: "三",
        3: "四",
        4: "五",
        5: "六",
        6: "日"
    }

    d = pd.to_datetime(date_str)

    return f"{d.month:02d}/{d.day:02d}({week_map[d.weekday()]})"


# =====================================
# 獎號格式
# =====================================

# 539、大樂透、威力彩使用
def format_lottery_numbers(nums):

    return "  ".join(
        f"{int(n):02d}"
        for n in nums
    )


# 3星彩、4星彩使用
def format_star_numbers(nums):

    return "  ".join(
        str(int(n))
        for n in nums
    )

# =====================================
# 複製列格式
# =====================================

def copy_row_style(ws, source_row, target_row):

    for col in range(1, ws.max_column + 1):

        src = ws.cell(source_row, col)
        dst = ws.cell(target_row, col)

        if src.has_style:

            dst.font = copy(src.font)
            dst.fill = copy(src.fill)
            dst.border = copy(src.border)
            dst.alignment = copy(src.alignment)
            dst.number_format = copy(src.number_format)
            dst.protection = copy(src.protection)

    ws.row_dimensions[target_row].height = \
        ws.row_dimensions[source_row].height


# =====================================
# 清除資料
# =====================================

def clear_rows(ws, start_row=2):

    for r in range(start_row, 200):

        for c in range(1, ws.max_column + 1):

            ws.cell(r, c).value = None


# =====================================
# 抓取資料
# =====================================


def get_latest_data(get_func, game_name):

    try:
        data = get_func()

        if len(data) > 0:
            print(f"{game_name}: 本月資料 {len(data)} 筆")
            return data[:10]

    except Exception as e:
        print(f"{game_name} 本月抓取失敗: {e}")

    print(f"{game_name}: 本月無資料，保留空資料")

    return []
# =====================================
# 抓取200期資料 
# =====================================
# 539
def get_539_history(target=200):
    crawler = TaiwanLotteryCrawler()
    result = []
    year = datetime.now().year
    month = datetime.now().month
    while len(result) < target:
        try:
            data = crawler.daily_cash([str(year), f"{month:02d}"])
            result.extend(data)
        except:
            pass
        month -= 1
        if month == 0:
            month = 12
            year -= 1
    return result[:target]
# 大樂透
def get_lotto649_history(target=200):

    crawler = TaiwanLotteryCrawler()

    result = []

    year = datetime.now().year
    month = datetime.now().month

    while len(result) < target:

        try:
            data = crawler.lotto649([str(year), f"{month:02d}"])
            result.extend(data)

        except:
            pass

        month -= 1

        if month == 0:
            month = 12
            year -= 1

    return result[:target]
#威力採
def get_power_history(target=200):

    crawler = TaiwanLotteryCrawler()

    result = []

    year = datetime.now().year
    month = datetime.now().month

    while len(result) < target:

        try:

            data = crawler.super_lotto(
                [str(year), f"{month:02d}"]
            )

            result.extend(data)

        except:
            pass

        month -= 1

        if month == 0:
            month = 12
            year -= 1

    return result[:target]
# =====================================
# 均值演算法
# =====================================

#539-----------------------------------
def predict_539_mean(records, periods):

    records = records[:periods]    #perids 使用者選幾筆資料切片

    counter = Counter()

    for item in records:

        for num in item["獎號"]:

            counter[num] += 1

    zones = [
        range(1, 9),
        range(9, 17),
        range(17, 25),
        range(25, 33),
        range(33, 40)
    ]

    result = []
    
    for zone in zones:
        nums = list(zone)
        avg = sum(counter.get(n, 0)for n in nums) / len(nums)
        best = min(nums,key=lambda n:abs(counter.get(n, 0) - avg))
        result.append(best)

    return sorted(result)

#大樂透-------------------------------------
def predict_lotto649_mean(records, periods):

    records = records[:periods]

    counter = Counter()

    for item in records:

        for num in item["獎號"]:

            counter[num] += 1

    zones = [
        range(1, 9),
        range(9, 17),
        range(17, 25),
        range(25, 33),
        range(33, 41),
        range(41, 50)
    ]

    result = []

    for zone in zones:

        nums = list(zone)

        avg = sum(counter.get(n, 0)for n in nums) / len(nums)
        best = min(nums,key=lambda n:abs(counter.get(n, 0) - avg))

        result.append(best)

    return sorted(result)
#威力彩-------------------------------------
def predict_power_mean(records, periods):

    records = records[:periods]

    first_counter = Counter()
    second_counter = Counter()

    for item in records:

        for num in item["第一區"]:
            first_counter[num] += 1

        second_counter[int(item["第二區"])] += 1

    zones = [
        range(1, 7),
        range(7, 13),
        range(13, 19),
        range(19, 25),
        range(25, 31),
        range(31, 39)
    ]

    first_result = []

    for zone in zones:
        nums = list(zone)
        avg = sum(first_counter.get(n, 0)for n in nums) / len(nums)
        best = min(nums,key=lambda n:abs(first_counter.get(n, 0) - avg))
        first_result.append(best)

    # 第二區
    second_nums = list(range(1, 9))
    avg = sum(second_counter.get(n, 0)for n in second_nums) / len(second_nums)
    second_best = min(second_nums,key=lambda n:abs(second_counter.get(n, 0) - avg))

    return {
        "first": sorted(first_result),
        "second": second_best
    }
# =====================================
# 抓取資料（自動補上月）
# =====================================
def get_data():
    
    crawler = TaiwanLotteryCrawler()

    now = datetime.now()

    cur_year = str(now.year)
    cur_month = f"{now.month:02d}"

    if now.month == 1:
        pre_year = str(now.year - 1)
        pre_month = "12"
    else:
        pre_year = str(now.year)
        pre_month = f"{now.month - 1:02d}"

    def merge_game_data(func):

        try:
            current = func([cur_year, cur_month])
        except Exception:
            current = []

        try:
            previous = func([pre_year, pre_month])
        except Exception:
            previous = []

        merged = current + previous

        seen = set()
        result = []

        for item in merged:

            issue = item["期別"]

            if issue not in seen:
                seen.add(issue)
                result.append(item)

        result.sort(key=lambda x: x["期別"],reverse=True)

        return result[:10]

    data = {
        "539": merge_game_data(crawler.daily_cash),
        "威力彩": merge_game_data(crawler.super_lotto),
        "大樂透": merge_game_data(crawler.lotto649),
        "3星彩": get_3star(),
        "4星彩": get_4star()
    }

    print("539:", len(data["539"]))
    print("威力彩:", len(data["威力彩"]))
    print("大樂透:", len(data["大樂透"]))
    print("3星彩:", len(data["3星彩"]))
    print("4星彩:", len(data["4星彩"]))

    return data
# =====================================
# 今彩539
# =====================================

def write_539(ws, data):

    clear_rows(ws)

    for r in range(3, 12):
        copy_row_style(ws, 2, r)

    for row, item in enumerate(reversed(data), start=2):

        ws.cell(row, 1).value = item["期別"]
        ws.cell(row, 2).value = roc_date(item["開獎日期"])
        ws.cell(row, 3).value = format_lottery_numbers(item["獎號"])


# =====================================
# 威力彩
# =====================================

def write_power(ws, data):

    clear_rows(ws)

    for r in range(3, 12):
        copy_row_style(ws, 2, r)

    for row, item in enumerate(reversed(data), start=2):

        ws.cell(row, 1).value = item["期別"]

        ws.cell(row, 2).value = roc_date(item["開獎日期"])

        ws.cell(row, 3).value = format_lottery_numbers(item["第一區"])

        ws.cell(row, 4).value = \
            f"{int(item['第二區']):02d}"


# =====================================
# 大樂透
# =====================================

def write_lotto649(ws, data):

    clear_rows(ws)

    for r in range(3, 12):
        copy_row_style(ws, 2, r)

    for row, item in enumerate(reversed(data), start=2):

        ws.cell(row, 1).value = item["期別"]
        ws.cell(row, 2).value = roc_date(item["開獎日期"])
        ws.cell(row, 3).value = format_lottery_numbers(item["獎號"])
        ws.cell(row, 4).value = \
            f"{int(item['特別號']):02d}"


# =====================================
# 3星彩
# =====================================

def write_3star(ws, data):

    clear_rows(ws)

    for r in range(3, 12):
        copy_row_style(ws, 2, r)

    for row, item in enumerate(reversed(data), start=2):

        ws.cell(row, 1).value = short_date(item["開獎日期"])
        ws.cell(row, 2).value = format_star_numbers(item["獎號"])


# =====================================
# 4星彩
# =====================================

def write_4star(ws, data):

    clear_rows(ws)

    for r in range(3, 12):
        copy_row_style(ws, 2, r)

    for row, item in enumerate(reversed(data), start=2):

        ws.cell(row, 1).value = short_date(item["開獎日期"])
        ws.cell(row, 2).value = format_star_numbers(item["獎號"])


# =====================================
# Excel
# =====================================

def create_excel(data):

    wb = load_workbook("excel範本.xlsx")

    write_539(wb["539"],data["539"])

    write_power(wb["威力彩"],data["威力彩"])

    write_lotto649(wb["大樂透"],data["大樂透"])

    write_3star(wb["3星"],data["3星彩"])

    write_4star(wb["4星"],data["4星彩"])
    wb.save("最新開獎紀錄.xlsx")

    print("Excel完成")

# =====================================
# Word 表格清空
# =====================================

def clear_table(table):

    for r in range(1, len(table.rows)):

        for cell in table.rows[r].cells:

            cell.text = ""


# =====================================
# Word 字型縮小
# =====================================

def resize_table_font(table):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                for run in paragraph.runs:

                    run.font.size = Pt(12.5)
                    # 字體大小
                    # 粗體
                    run.font.bold = True

                    # 英文、數字
                    run.font.name = "Franklin Gothic Heavy"

                    # 中文
                    run._element.rPr.rFonts.set(qn("w:eastAsia"),"DFKai-SB")

# =====================================
# Word產生
# =====================================

def create_word(data):

    doc = Document("word 範本.docx")

    tables = doc.tables

    # 539
    table = tables[0]

    clear_table(table)

    for i, item in enumerate(reversed(data["539"])):

        row = i + 1

        table.cell(row, 0).text = str(item["期別"])
        table.cell(row, 1).text = roc_date(item["開獎日期"])
        table.cell(row, 2).text = format_lottery_numbers(item["獎號"])

    resize_table_font(table)

    # 3星彩
    table = tables[1]

    clear_table(table)

    for i, item in enumerate(reversed(data["3星彩"])):

        row = i + 1
        table.cell(row, 0).text = short_date(item["開獎日期"])
        table.cell(row, 1).text = format_star_numbers(item["獎號"])

    resize_table_font(table)

    # 威力彩
    table = tables[2]

    clear_table(table)

    for i, item in enumerate(reversed(data["威力彩"])):

        row = i + 1

        table.cell(row, 0).text = str(item["期別"])
        table.cell(row, 1).text = roc_date(item["開獎日期"])
        table.cell(row, 2).text = format_lottery_numbers(item["第一區"])
        table.cell(row, 3).text = \
            f"{int(item['第二區']):02d}"

    resize_table_font(table)

    # 4星彩
    table = tables[3]

    clear_table(table)

    for i, item in enumerate(reversed(data["4星彩"])):
        row = i + 1
        table.cell(row, 0).text = short_date(item["開獎日期"])
        table.cell(row, 1).text = format_star_numbers(item["獎號"])

    resize_table_font(table)

    # 大樂透
    table = tables[4]

    clear_table(table)

    for i, item in enumerate(reversed(data["大樂透"])):
        row = i + 1
        table.cell(row, 0).text = str(item["期別"])
        table.cell(row, 1).text = roc_date(item["開獎日期"])
        table.cell(row, 2).text = format_lottery_numbers(item["獎號"])
        table.cell(row, 3).text = \
            f"{int(item['特別號']):02d}"

    resize_table_font(table)

    doc.save("最新開獎紀錄.docx")

    print("Word完成")
# ====================================
#熱度區間
# ====================================
def get_539_heat(records, periods):
    records = records[:periods]
    zones = [
        range(1, 9),
        range(9, 17),
        range(17, 25),
        range(25, 33),
        range(33, 40)
    ]
    heat = []
    for zone in zones:
        count = 0
        for item in records:
            for num in item["獎號"]:
                if num in zone:
                    count += 1
        heat.append(count)

    return heat

def get_lotto649_heat(records, periods):
    records = records[:periods]
    zones = [
        range(1, 9),
        range(9, 17),
        range(17, 25),
        range(25, 33),
        range(33, 41),
        range(41, 50)
    ]
    heat = []
    for zone in zones:
        count = 0
        for item in records:
            for num in item["獎號"]:
                if num in zone:
                    count += 1
        heat.append(count)

    return heat

def get_power_heat(records, periods):

    records = records[:periods]

    zones = [
        range(1, 7),
        range(7, 13),
        range(13, 19),
        range(19, 25),
        range(25, 31),
        range(31, 39)
    ]

    heat = []

    for zone in zones:

        count = 0

        for item in records:

            for num in item["第一區"]:

                if num in zone:
                    count += 1

        heat.append(count)

    second_counter = Counter()

    for item in records:

        second_counter[int(item["第二區"])] += 1

    second_heat = []

    for n in range(1, 9):

        second_heat.append(
            second_counter.get(n, 0)
        )

    return {
        "zones": heat,
        "second": second_heat
    }

# =====================================
# 產生json
# =====================================
#顯示開獎資料
def create_json(data):
    if (
        not data["539"]
        or not data["威力彩"]
        or not data["大樂透"]
        or not data["3星彩"]
        or not data["4星彩"]
    ):
        print("部分資料不存在，略過 JSON")
        return

    result = {

        "update_time": str(pd.Timestamp.now()),

        "539": {
            "period": str(
                data["539"][0]["期別"]
            ),

            "date": roc_date(
                data["539"][0]["開獎日期"]
            ),

            "numbers": format_lottery_numbers(
                data["539"][0]["獎號"]
            )
        },

        "power": {

            "period": str(
                data["威力彩"][0]["期別"]
            ),

            "date": roc_date(
                data["威力彩"][0]["開獎日期"]
            ),

            "numbers": format_lottery_numbers(
                data["威力彩"][0]["第一區"]
            ),

            "special": f"{int(data['威力彩'][0]['第二區']):02d}"
        },

        "lotto649": {

            "period": str(
                data["大樂透"][0]["期別"]
            ),

            "date": roc_date(
                data["大樂透"][0]["開獎日期"]
            ),

            "numbers": format_lottery_numbers(
                data["大樂透"][0]["獎號"]
            ),

            "special": f"{int(data['大樂透'][0]['特別號']):02d}"
        },

        "3star": {

            "period": str(
                data["3星彩"][0]["期別"]
            ),

            "date": short_date(
                data["3星彩"][0]["開獎日期"]
            ),

            "numbers": format_star_numbers(
                data["3星彩"][0]["獎號"]
            )
        },

        "4star": {

            "period": str(
                data["4星彩"][0]["期別"]
            ),

            "date": short_date(
                data["4星彩"][0]["開獎日期"]
            ),

            "numbers": format_star_numbers(
                data["4星彩"][0]["獎號"]
            )
        }
    }

    with open(
        "data.json",
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            result,
            f,
            ensure_ascii=False,
            indent=4
        )

    print("data.json完成")
#-------------------------------------   
# 預測json
def create_prediction_json():

    history539 = get_539_history(200)
    history649 = get_lotto649_history(200)
    historyPower = get_power_history(200)
    prediction = {

        "539": {
            "10": {"numbers": predict_539_mean(history539,10),
                   "heat": get_539_heat(history539,10)},                
            "30":{"numbers": predict_539_mean(history539,30),
                   "heat": get_539_heat(history539,30)},
            "50":{"numbers": predict_539_mean(history539,50),
                   "heat": get_539_heat(history539,50)},
            "100":{"numbers": predict_539_mean(history539,100),
                   "heat": get_539_heat(history539,100)},
            "200":{"numbers": predict_539_mean(history539,200),
                   "heat": get_539_heat(history539,200)},
        },

        "lotto649": {
            "10": {"numbers": predict_lotto649_mean(history649,10),
                   "heat": get_lotto649_heat(history649,10)},
            "30": {"numbers": predict_lotto649_mean(history649,30),
                   "heat": get_lotto649_heat(history649,30)},
            "50": {"numbers": predict_lotto649_mean(history649,50),
                   "heat": get_lotto649_heat(history649,50)},
            "100": {"numbers": predict_lotto649_mean(history649,100),
                   "heat": get_lotto649_heat(history649,100)},
            "200": {"numbers": predict_lotto649_mean(history649,200),
                   "heat": get_lotto649_heat(history649,200)},
        },

        "power": {

            "10": {"prediction": predict_power_mean(historyPower,10),
                   "heat": get_power_heat(historyPower,10)},
            "30": {"prediction": predict_power_mean(historyPower,30),
                   "heat": get_power_heat(historyPower,30)},
            "50": {"prediction": predict_power_mean(historyPower,50),
                   "heat": get_power_heat(historyPower,50)},
            "100": {"prediction": predict_power_mean(historyPower,100),
                   "heat": get_power_heat(historyPower,100)},
            "200": {"prediction": predict_power_mean(historyPower,200),
                   "heat": get_power_heat(historyPower,200)}
        
}

    }

    with open(
        "prediction.json",
        "w",
        encoding="utf-8"
    ) as f:

        json.dump(
            prediction,
            f,
            ensure_ascii=False,
            indent=4
        )

    print("prediction.json完成")

# =====================================
# 主程式
# =====================================

def main():

    print("開始抓取資料...")

    data = get_data()

    if all(len(v) == 0 for v in data.values()):

        print("查無任何資料")
        return

    create_excel(data)
    create_word(data)
    create_json(data)
    create_prediction_json()
    create_scratch_json()
    print("全部完成")

    
if __name__ == "__main__":
    main()
